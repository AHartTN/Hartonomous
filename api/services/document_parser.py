"""
Document Parser Service
Atomizes PDF, DOCX, Markdown, HTML, TXT into extreme granular atoms.

Hierarchical Pattern:
  document → sections/pages → paragraphs → sentences → words → characters

Every level becomes an atom with metadata and hierarchical composition.

Copyright (c) 2025 Anthony Hart. All Rights Reserved.
"""

import hashlib
import io
import json
import logging
from pathlib import Path
from typing import Any, Dict, Optional

from psycopg import AsyncConnection

logger = logging.getLogger(__name__)


class DocumentParserService:
    """Service for parsing and atomizing documents."""

    @staticmethod
    async def parse_and_atomize_pdf(
        conn: AsyncConnection,
        file_path: Optional[Path] = None,
        file_data: Optional[bytes] = None,
        metadata: Optional[Dict[str, Any]] = None,
        extract_images: bool = True,
        ocr_enabled: bool = False,
    ) -> Dict[str, Any]:
        """
        Parse and atomize PDF document.

        Args:
            conn: Database connection
            file_path: Path to PDF file (or file_data)
            file_data: PDF file bytes
            metadata: Additional metadata
            extract_images: Extract and atomize embedded images
            ocr_enabled: Use OCR for scanned PDFs

        Returns:
            dict with atom counts, root_atom_id, page_ids
        """
        try:
            import pdfplumber

            if metadata is None:
                metadata = {}

            # Open PDF
            if file_data:
                pdf_file = pdfplumber.open(io.BytesIO(file_data))
            elif file_path:
                pdf_file = pdfplumber.open(file_path)
            else:
                raise ValueError("Either file_path or file_data must be provided")

            with pdf_file as pdf:
                # Extract PDF metadata
                pdf_metadata = pdf.metadata or {}
                metadata.update(
                    {
                        "modality": "document",
                        "format": "pdf",
                        "total_pages": len(pdf.pages),
                        "pdf_metadata": pdf_metadata,
                    }
                )

                # Create document root atom
                doc_title = pdf_metadata.get(
                    "Title", file_path.name if file_path else "document"
                )

                async with conn.cursor() as cur:
                    # Create document atom
                    await cur.execute(
                        """
                        SELECT atomize_value(
                            digest(%s, 'sha256'),
                            %s,
                            %s::jsonb
                        )
                        """,
                        (doc_title.encode("utf-8"), doc_title, json.dumps(metadata)),
                    )
                    doc_atom_id = (await cur.fetchone())[0]

                    page_ids = []
                    total_atoms = 1  # Document atom

                    # Process each page
                    for page_num, page in enumerate(pdf.pages, start=1):
                        # Extract text
                        text = page.extract_text() or ""

                        # OCR if needed and enabled
                        if ocr_enabled and not text.strip():
                            text = await DocumentParserService._ocr_page(page)

                        # Create page atom
                        page_metadata = {
                            "modality": "page",
                            "page_number": page_num,
                            "width": page.width,
                            "height": page.height,
                        }

                        await cur.execute(
                            """
                            SELECT atomize_value(
                                digest(%s, 'sha256'),
                                %s,
                                %s::jsonb
                            )
                            """,
                            (
                                f"page_{page_num}".encode("utf-8"),
                                f"Page {page_num}",
                                json.dumps(page_metadata),
                            ),
                        )
                        page_atom_id = (await cur.fetchone())[0]
                        page_ids.append(page_atom_id)
                        total_atoms += 1

                        # Link page to document using SQL function
                        await cur.execute(
                            """
                            SELECT create_composition(
                                %s::bigint,
                                %s::bigint,
                                %s::bigint,
                                '{}'::jsonb
                            )
                            """,
                            (doc_atom_id, page_atom_id, page_num - 1),
                        )

                        # Atomize page text
                        if text.strip():
                            await cur.execute(
                                "SELECT atomize_text(%s, %s::jsonb)",
                                (text, json.dumps({"page": page_num})),
                            )
                            char_atoms = (await cur.fetchone())[0]
                            total_atoms += len(char_atoms)

                            # Link text atoms to page using SQL function
                            for idx, char_atom_id in enumerate(char_atoms):
                                await cur.execute(
                                    """
                                    SELECT create_composition(
                                        %s::bigint,
                                        %s::bigint,
                                        %s::bigint,
                                        '{}'::jsonb
                                    )
                                    """,
                                    (page_atom_id, char_atom_id, idx),
                                )

                        # Extract and atomize images
                        if extract_images and page.images:
                            for img_idx, img in enumerate(page.images):
                                logger.info(
                                    f"Image on page {page_num}: "
                                    f"{img.get('width', 'unknown')}x{img.get('height', 'unknown')}"
                                )

                                # Create image atom with metadata
                                # In future: extract pixels, vectorize, create spatial atoms
                                img_metadata = {
                                    "modality": "image",
                                    "format": img.get("ext", "unknown"),
                                    "page": page_num,
                                    "index": img_idx,
                                    "width": img.get("width"),
                                    "height": img.get("height"),
                                }

                                img_ref = f"image_{page_num}_{img_idx}"
                                await cur.execute(
                                    """
                                    SELECT atomize_value(
                                        digest(%s, 'sha256'),
                                        %s,
                                        %s::jsonb
                                    )
                                    """,
                                    (
                                        img_ref.encode("utf-8"),
                                        img_ref,
                                        json.dumps(img_metadata),
                                    ),
                                )
                                img_atom_id = (await cur.fetchone())[0]
                                total_atoms += 1

                                # Link image to page
                                await cur.execute(
                                    """
                                    SELECT create_composition(
                                        %s::bigint,
                                        %s::bigint,
                                        %s::bigint,
                                        jsonb_build_object('type', 'contains_image')
                                    )
                                    """,
                                    (page_atom_id, img_atom_id, img_idx),
                                )

                    logger.info(
                        f"PDF atomization complete: {len(page_ids)} pages, "
                        f"{total_atoms} total atoms"
                    )

                    return {
                        "atom_count": total_atoms,
                        "root_atom_id": doc_atom_id,
                        "page_ids": page_ids,
                        "total_pages": len(pdf.pages),
                    }

        except Exception as e:
            logger.error(f"PDF parsing failed: {e}", exc_info=True)
            raise

    @staticmethod
    async def parse_and_atomize_docx(
        conn: AsyncConnection,
        file_path: Optional[Path] = None,
        file_data: Optional[bytes] = None,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """
        Parse and atomize DOCX document.

        Structure: document → sections → paragraphs → runs → words → characters
        """
        try:
            from docx import Document

            if metadata is None:
                metadata = {}

            # Open DOCX
            if file_data:
                doc = Document(io.BytesIO(file_data))
            elif file_path:
                doc = Document(file_path)
            else:
                raise ValueError("Either file_path or file_data must be provided")

            # Extract metadata
            props = doc.core_properties
            metadata.update(
                {
                    "modality": "document",
                    "format": "docx",
                    "title": props.title
                    or (file_path.name if file_path else "document"),
                    "author": props.author,
                    "created": str(props.created) if props.created else None,
                }
            )

            async with conn.cursor() as cur:
                # Create document atom
                doc_title = props.title or (file_path.name if file_path else "document")

                await cur.execute(
                    """
                    SELECT atomize_value(
                        digest(%s, 'sha256'),
                        %s,
                        %s::jsonb
                    )
                    """,
                    (doc_title.encode("utf-8"), doc_title, json.dumps(metadata)),
                )
                doc_atom_id = (await cur.fetchone())[0]

                total_atoms = 1
                para_ids = []

                # Process paragraphs
                for para_idx, para in enumerate(doc.paragraphs):
                    if not para.text.strip():
                        continue

                    # Create paragraph atom
                    para_metadata = {
                        "modality": "paragraph",
                        "style": para.style.name,
                        "alignment": str(para.alignment) if para.alignment else None,
                    }

                    await cur.execute(
                        """
                        SELECT atomize_value(
                            digest(%s, 'sha256'),
                            %s,
                            %s::jsonb
                        )
                        """,
                        (
                            para.text[:64].encode("utf-8"),
                            para.text[:100],
                            json.dumps(para_metadata),
                        ),
                    )
                    para_atom_id = (await cur.fetchone())[0]
                    para_ids.append(para_atom_id)
                    total_atoms += 1

                    # Link paragraph to document using SQL function
                    await cur.execute(
                        """
                        SELECT create_composition(
                            %s::bigint,
                            %s::bigint,
                            %s::bigint,
                            '{}'::jsonb
                        )
                        """,
                        (doc_atom_id, para_atom_id, para_idx),
                    )

                    # Atomize paragraph text
                    await cur.execute(
                        "SELECT atomize_text(%s, %s::jsonb)",
                        (para.text, json.dumps({"paragraph": para_idx})),
                    )
                    char_atoms = (await cur.fetchone())[0]
                    total_atoms += len(char_atoms)

                    # Link text atoms to paragraph using SQL function
                    for idx, char_atom_id in enumerate(char_atoms):
                        await cur.execute(
                            """
                            SELECT create_composition(
                                %s::bigint,
                                %s::bigint,
                                %s::bigint,
                                '{}'::jsonb
                            )
                            """,
                            (para_atom_id, char_atom_id, idx),
                        )

                # Process tables
                for table_idx, table in enumerate(doc.tables):
                    logger.info(
                        f"Table {table_idx}: {len(table.rows)} rows x {len(table.columns)} columns"
                    )

                    # Create table atom
                    table_ref = f"table_{table_idx}"
                    table_metadata = {
                        "modality": "table",
                        "rows": len(table.rows),
                        "columns": len(table.columns),
                        "index": table_idx,
                    }

                    await cur.execute(
                        """
                        SELECT atomize_value(
                            digest(%s, 'sha256'),
                            %s,
                            %s::jsonb
                        )
                        """,
                        (
                            table_ref.encode("utf-8"),
                            table_ref,
                            json.dumps(table_metadata),
                        ),
                    )
                    table_atom_id = (await cur.fetchone())[0]
                    total_atoms += 1

                    # Link table to document
                    await cur.execute(
                        """
                        SELECT create_composition(
                            %s::bigint,
                            %s::bigint,
                            %s::bigint,
                            jsonb_build_object('type', 'contains_table')
                        )
                        """,
                        (doc_atom_id, table_atom_id, table_idx),
                    )

                    # Atomize each cell
                    for row_idx, row in enumerate(table.rows):
                        for col_idx, cell in enumerate(row.cells):
                            cell_text = cell.text.strip()
                            if cell_text:
                                cell_metadata = {
                                    "modality": "table_cell",
                                    "row": row_idx,
                                    "column": col_idx,
                                    "table_index": table_idx,
                                }

                                await cur.execute(
                                    "SELECT atomize_text(%s, %s::jsonb)",
                                    (cell_text, json.dumps(cell_metadata)),
                                )
                                cell_char_atoms = (await cur.fetchone())[0]

                                # Create cell atom and link to table
                                cell_hash = hashlib.sha256(cell_text.encode()).digest()
                                await cur.execute(
                                    """
                                    SELECT atomize_value(%s, %s, %s::jsonb)
                                    """,
                                    (cell_hash, cell_text, json.dumps(cell_metadata)),
                                )
                                cell_atom_id = (await cur.fetchone())[0]
                                total_atoms += len(cell_char_atoms) + 1

                                # Link cell to table with position
                                cell_position = row_idx * len(table.columns) + col_idx
                                await cur.execute(
                                    """
                                    SELECT create_composition(
                                        %s::bigint,
                                        %s::bigint,
                                        %s::bigint,
                                        jsonb_build_object('row', %s, 'col', %s)
                                    )
                                    """,
                                    (
                                        table_atom_id,
                                        cell_atom_id,
                                        cell_position,
                                        row_idx,
                                        col_idx,
                                    ),
                                )

                logger.info(
                    f"DOCX atomization complete: {len(para_ids)} paragraphs, "
                    f"{total_atoms} total atoms"
                )

                return {
                    "atom_count": total_atoms,
                    "root_atom_id": doc_atom_id,
                    "paragraph_ids": para_ids,
                }

        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}", exc_info=True)
            raise

    @staticmethod
    async def parse_and_atomize_markdown(
        conn: AsyncConnection,
        text: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """
        Parse and atomize Markdown document.

        Structure: document → sections (by headers) → blocks → text
        """
        try:
            from markdown_it import MarkdownIt

            if metadata is None:
                metadata = {}

            md = MarkdownIt()
            tokens = md.parse(text)

            metadata.update({"modality": "document", "format": "markdown"})

            # Extract title from first h1 heading
            title = "Markdown Document"  # default
            for i, token in enumerate(tokens):
                if token.type == "heading_open" and token.tag == "h1":
                    # Next token should be inline with the heading text
                    if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                        title = tokens[i + 1].content.strip()
                        break

            async with conn.cursor() as cur:
                # Create document atom
                await cur.execute(
                    """
                    SELECT atomize_value(
                        digest(%s, 'sha256'),
                        %s,
                        %s::jsonb
                    )
                    """,
                    (title.encode("utf-8"), title, json.dumps(metadata)),
                )
                doc_atom_id = (await cur.fetchone())[0]

                total_atoms = 1

                # Process markdown tokens
                for token in tokens:
                    if token.type == "heading_open":
                        # Create section atom for heading
                        pass
                    elif token.type == "paragraph_open":
                        # Create paragraph atom
                        pass
                    elif token.type == "code_block" or token.type == "fence":
                        # Create code block atom
                        code = token.content
                        lang = token.info if hasattr(token, "info") else "plaintext"

                        code_metadata = metadata.copy()
                        code_metadata.update(
                            {
                                "modality": "code",
                                "language": lang,
                            }
                        )

                        # C# code: use semantic AST atomizer for deep code understanding
                        if lang.lower() in ("csharp", "cs", "c#"):
                            try:
                                from api.services.code_atomization.code_atomizer_client import (
                                    CodeAtomizerClient,
                                )

                                client = CodeAtomizerClient()
                                try:
                                    # Check if C# microservice is available
                                    if await client.health_check():
                                        logger.info(
                                            f"Atomizing C# code via Roslyn microservice (AST-level)..."
                                        )

                                        # Get AST from C# microservice
                                        ast_result = await client.atomize_csharp(
                                            code=code,
                                            filename="markdown_code_block.cs",
                                            metadata=json.dumps(code_metadata),
                                        )

                                        # Create code block atom with AST metadata
                                        code_hash = hashlib.sha256(
                                            code.encode()
                                        ).digest()
                                        code_metadata["ast_available"] = True
                                        code_metadata["ast_nodes"] = ast_result.get(
                                            "node_count", 0
                                        )

                                        await cur.execute(
                                            """
                                            SELECT atomize_value(
                                                %s,
                                                %s,
                                                %s::jsonb
                                            )
                                            """,
                                            (
                                                code_hash,
                                                code[:100],
                                                json.dumps(code_metadata),
                                            ),
                                        )
                                        (await cur.fetchone())[0]
                                        total_atoms += 1

                                        logger.info(
                                            f"✓ C# code atomized via Roslyn AST ({ast_result.get('node_count', 0)} nodes)"
                                        )
                                    else:
                                        logger.warning(
                                            "C# atomizer service unavailable, falling back to text atomization"
                                        )
                                        raise RuntimeError(
                                            "C# atomizer service unavailable"
                                        )

                                finally:
                                    await client.close()

                            except Exception as e:
                                logger.warning(
                                    f"C# atomization failed: {e}, falling back to text atomization"
                                )
                                # Fallback to text atomization
                                await cur.execute(
                                    "SELECT atomize_text(%s, %s::jsonb)",
                                    (code, json.dumps(code_metadata)),
                                )
                                code_atoms = (await cur.fetchone())[0]
                                total_atoms += len(code_atoms)
                        else:
                            # Other languages: atomize as text
                            await cur.execute(
                                "SELECT atomize_text(%s, %s::jsonb)",
                                (code, json.dumps(code_metadata)),
                            )
                            code_atoms = (await cur.fetchone())[0]
                            total_atoms += len(code_atoms)
                    elif token.type == "inline":
                        # Atomize inline text
                        if token.content:
                            await cur.execute(
                                "SELECT atomize_text(%s, %s::jsonb)",
                                (token.content, json.dumps(metadata)),
                            )
                            char_atoms = (await cur.fetchone())[0]
                            total_atoms += len(char_atoms)

                return {"atom_count": total_atoms, "root_atom_id": doc_atom_id}

        except Exception as e:
            logger.error(f"Markdown parsing failed: {e}", exc_info=True)
            raise

    @staticmethod
    async def _ocr_page(page) -> str:
        """
        Perform OCR on PDF page using Tesseract.
        """
        try:
            import pytesseract

            # Convert page to image
            img = page.to_image(resolution=300)
            pil_img = img.original

            # Perform OCR
            text = pytesseract.image_to_string(pil_img, lang="eng")

            logger.info(f"OCR extracted {len(text)} characters")
            return text

        except Exception as e:
            logger.error(f"OCR failed: {e}")
            return ""


__all__ = ["DocumentParserService"]
